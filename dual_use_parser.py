"""
Dual-Use List Parser
Parses the dual-use section using style IDs to determine hierarchy
"""

import json
import re
from pathlib import Path
from typing import List, Dict, Any, Optional
from docx import Document


class DualUseParser:
    """Parser for the Dual-Use List section using style IDs"""

    def __init__(self):
        # Define style hierarchy for content items
        self.content_styles = {
            'ActHead 3': 0,  # Category headers
            'Cat0 new style 1': 1,  # Main section headers (A, B, C, D, E)
            'Cat0 new style 2': 2,  # Item codes like 0.A.001, 0.B.002
            'Cat0 new style 3 (1a1) paragraphs indented': 2,  # 1C241, 1C350, etc.
            'Cat0 new style 4 (1a1a) paragraphs indented': 3,  # 1C351.b, etc.
            'Cat0 new style 5 (1a1a1) paragraphs indented': 4,  # 1C351.a.17, etc.
            'Cat0 new style 6 (1a1a1a) paragraphs indented': 5,  # Fourth level
            'Cat0 new style 7 (1a1a1a1) paragraphs indented': 6,  # Fifth level
        }

        # Note styles - these should be treated differently
        self.note_styles = {
            'DL0ANote', 'DL0ANote(a)', 'DL0ANote1(a)', 'DL0Aa1Note',
            'DL0AaNote', 'DL0AaNote(a)', 'DL0AaNote1', 'DL0AaNotea1',
            'DL0ANoteRcN', 'DL0ATechH', 'DL0ATechText1'
        }

        # N.B. reference styles
        self.nb_styles = {'DL0ANB', 'DL0AaNB', 'DL0aNB'}

        # Item styles (lettered/numbered items within entries)
        # These map to specific hierarchy levels
        self.item_style_levels = {
            'DL0Aa': 3,      # Letter items (a., b., c.)
            'DL0Aa1': 4,     # Number items under letters (1., 2., 3.)
            'DL0Aa1a': 5,    # Letter items under numbers (a., b., c.)
            'DL0Aa1a1': 6,   # Number items under letters (1., 2., 3.)
        }

        # Patterns for parsing item labels
        self.patterns = {
            'category': re.compile(r'^Category\s+(\d+)[—\-\u2014]\s*(.*)$'),
            'section': re.compile(r'^(\d+)\.\s+([A-E])\.\s+(.*)$'),
            'item_code': re.compile(r'^(\d+)\.\s+([A-E])\.\s+(\d+)\.\s+(.*)$'),
            'item_code_alt': re.compile(r'^(\d+)\.\s+([A-E])\.\s*\t\s*(\d+)\.\s+(.*)$'),  # Format: 0. A. [tab] 001. ...
            'sub_item': re.compile(r'^(\d+)\.\s+([A-E])\.\s+(\d+)\.\s+([a-z])\.\s+(.*)$'),
            'letter_item': re.compile(r'^([a-z])\.\s+(.*)$'),
            'number_item': re.compile(r'^(\d+)\.\s+(.*)$'),
        }

    def find_dual_use_section(self, doc: Document) -> tuple[Optional[int], Optional[int]]:
        """Find the start and end of the dual-use section"""
        start_idx = None
        end_idx = None

        for i, para in enumerate(doc.paragraphs):
            text = para.text.strip()
            style_name = para.style.name if para.style else "Normal"

            # Look for start - "Category 0" with ActHead style (not TOC)
            if start_idx is None:
                if ('Category 0' in text and
                    'Nuclear materials' in text and
                    'ActHead' in style_name):
                    start_idx = i
                    print(f"Found dual-use section start at paragraph {i}: {text}")

            # Look for end
            if start_idx and end_idx is None:
                if 'Schedule' in text and 'Repeal' in text:
                    end_idx = i
                    print(f"Found dual-use section end at paragraph {i}")
                    break

        return start_idx, end_idx

    def get_style_level(self, style_name: str) -> Optional[int]:
        """Get the hierarchical level for a content style"""
        return self.content_styles.get(style_name)

    def is_note_style(self, style_name: str) -> bool:
        """Check if this is a note style"""
        return style_name in self.note_styles or style_name in self.nb_styles

    def get_item_style_level(self, style_name: str) -> Optional[int]:
        """Get the level for an item style, or None if not an item style"""
        return self.item_style_levels.get(style_name)

    def parse_label(self, text: str, style_name: str) -> Dict[str, Any]:
        """Parse the label and description from text based on patterns"""
        text = text.strip()

        # Check if this is a Note - should be filtered out
        if text.startswith('Note:') or text.startswith('Note\t'):
            return {
                'type': 'note',
                'label': '',
                'description': text
            }

        # Try category pattern
        match = self.patterns['category'].match(text)
        if match:
            return {
                'type': 'category',
                'label': f"Category {match.group(1)}",
                'description': match.group(2).strip()
            }

        # Try section pattern (1. A. Software)
        match = self.patterns['section'].match(text)
        if match:
            return {
                'type': 'section',
                'label': f"{match.group(1)}.{match.group(2)}",
                'description': match.group(3).strip()
            }

        # Try item code pattern (1. C. 350. ...)
        match = self.patterns['item_code'].match(text)
        if match:
            return {
                'type': 'item_code',
                'label': f"{match.group(1)}{match.group(2)}{match.group(3)}",
                'description': match.group(4).strip()
            }

        # Try alternative item code pattern (0. A. [tab] 001. ...)
        match = self.patterns['item_code_alt'].match(text)
        if match:
            return {
                'type': 'item_code',
                'label': f"{match.group(1)}{match.group(2)}{match.group(3)}",
                'description': match.group(4).strip()
            }

        # Try sub-item pattern (1. C. 351. b. ...)
        match = self.patterns['sub_item'].match(text)
        if match:
            return {
                'type': 'sub_item',
                'label': f"{match.group(1)}{match.group(2)}{match.group(3)}.{match.group(4)}",
                'description': match.group(5).strip()
            }

        # Try letter item (a., b., c.)
        match = self.patterns['letter_item'].match(text)
        if match:
            return {
                'type': 'letter',
                'label': match.group(1) + '.',
                'description': match.group(2).strip()
            }

        # Try number item (1., 2., 3.)
        match = self.patterns['number_item'].match(text)
        if match:
            return {
                'type': 'number',
                'label': match.group(1) + '.',
                'description': match.group(2).strip()
            }

        # Default - use full text as description
        return {
            'type': 'unknown',
            'label': '',
            'description': text
        }

    def build_hierarchy(self, paragraphs: List, start_idx: int, end_idx: int) -> List[Dict[str, Any]]:
        """Build hierarchical structure using style information"""
        result = []
        stack = []  # Stack of (node, level) tuples

        current_notes = []  # Accumulate notes for the next item

        for i in range(start_idx, end_idx):
            para = paragraphs[i]
            text = para.text.strip()

            if not text:
                continue

            style_name = para.style.name if para.style else "Normal"

            # Check if this is a content style
            level = self.get_style_level(style_name)

            # If not a Cat0 style, check if it's an item style (DL0Aa, DL0Aa1, etc.)
            if level is None:
                level = self.get_item_style_level(style_name)

            if level is not None:
                # This is a hierarchical content item
                parsed = self.parse_label(text, style_name)

                # Skip notes - don't add them to the hierarchy
                if parsed['type'] == 'note':
                    continue

                # Skip unlabeled continuation text (likely part of a note)
                # These are paragraphs with item styles but no proper label pattern
                if parsed['type'] == 'unknown' and not parsed['label']:
                    # Check if it looks like a continuation (doesn't start with a label)
                    if not re.match(r'^(\d+\.|\d+\.\s*[A-E]\.?|\d+\.\s*[a-z]\.?|[a-z]\.)', text):
                        continue

                # Clean up description and extract proper label
                description = parsed['description']
                label = parsed['label']

                # Check if label is a generic section label (like "0.A", "1.B", etc.)
                is_generic_label = label and re.match(r'^\d+\.[A-E]$', label)

                if is_generic_label:
                    # Try to extract the actual label from the description
                    # Pattern 1: "001.\tb.\t4.\t..." -> label="4.", description="..."
                    extract_match = re.match(r'^(\d+)\.\s*([a-z])\.\s*(\d+)\.\s*(.*)$', description)
                    if extract_match:
                        label = extract_match.group(3) + '.'
                        description = extract_match.group(4)
                    else:
                        # Pattern 2: "001.\tb.\t4.\ta.\t..." -> label="4.", description="a.\t..."
                        extract_match = re.match(r'^(\d+)\.\s*([a-z])\.\s*(\d+)\.\s*([a-z])\.\s*(.*)$', description)
                        if extract_match:
                            label = extract_match.group(3) + '.'
                            description = extract_match.group(4) + '.\t' + extract_match.group(5)
                        else:
                            # Pattern 3: "001.\tj.\t..." -> label="j.", description="..."
                            extract_match = re.match(r'^(\d+)\.\s*([a-z])\.\s*(.*)$', description)
                            if extract_match:
                                label = extract_match.group(2) + '.'
                                description = extract_match.group(3)
                            else:
                                # Pattern 4: "001.\t..." -> label="001.", description="..."
                                extract_match = re.match(r'^(\d+)\.\s*(.*)$', description)
                                if extract_match:
                                    label = extract_match.group(1) + '.'
                                    description = extract_match.group(2)

                node = {
                    'Label': label if label else text[:50],
                    'Description': description,
                    'SubStructures': []
                }

                # Clear accumulated notes (we're omitting them in output)
                current_notes = []

                # Determine effective level for hierarchy placement
                # If this is a letter item (a., b., c.) and the previous item at the same level
                # is a number item (1., 2., 13., etc.), treat this as a child of that item
                effective_level = level
                is_letter_label = re.match(r'^[a-z]\.$', label)

                if is_letter_label and stack:
                    # Check if previous item at same level has a number label OR is another letter
                    prev_node, prev_level = stack[-1]
                    prev_label = prev_node.get('Label', '')
                    is_prev_number = re.match(r'^\d+\.$', prev_label)
                    is_prev_letter = re.match(r'^[a-z]\.$', prev_label)

                    if prev_level == level and is_prev_number:
                        # Previous item is a number at same level - nest under it
                        effective_level = level + 1
                    elif prev_level == level + 1 and is_prev_letter:
                        # Previous item is a letter at level+1 - we're continuing the sequence
                        effective_level = level + 1

                # Find correct parent in stack based on effective level
                while stack and stack[-1][1] >= effective_level:
                    stack.pop()

                if stack:
                    # Add to parent's SubStructures
                    stack[-1][0]['SubStructures'].append(node)
                else:
                    # Top-level item
                    result.append(node)

                # Push current node onto stack with effective level
                stack.append((node, effective_level))

            elif self.is_note_style(style_name):
                # This is a note - accumulate it
                current_notes.append({
                    'Text': text,
                    'StyleID': style_name
                })

            else:
                # Unknown style - accumulate as note if substantial
                if text and len(text) > 10:  # Only substantial text
                    current_notes.append({
                        'Text': text,
                        'StyleID': style_name
                    })

        return result

    def parse_document(self, file_path: str) -> List[Dict[str, Any]]:
        """Parse the dual-use list from a DOCX file"""
        print(f"Loading document: {file_path}")
        doc = Document(file_path)
        print(f"Loaded {len(doc.paragraphs)} paragraphs")

        # Find dual-use section
        start_idx, end_idx = self.find_dual_use_section(doc)

        if start_idx is None:
            print("Error: Could not find dual-use section start")
            return []

        if end_idx is None:
            end_idx = len(doc.paragraphs)
            print("Using end of document as section end")

        print(f"\nParsing dual-use section: paragraphs {start_idx} to {end_idx}")

        # Build hierarchy
        hierarchy = self.build_hierarchy(doc.paragraphs, start_idx, end_idx)

        print(f"Created {len(hierarchy)} top-level items")

        return hierarchy


def count_items(structures: List[Dict[str, Any]]) -> int:
    """Recursively count all items in the hierarchy"""
    count = len(structures)
    for item in structures:
        count += count_items(item.get('SubStructures', []))
    return count


def print_hierarchy_sample(structures: List[Dict[str, Any]], depth: int = 0, max_depth: int = 3):
    """Print a sample of the hierarchy structure"""
    if depth > max_depth:
        return

    indent = "  " * depth
    for i, item in enumerate(structures[:5]):  # Show first 5 at each level
        label = item.get('Label', 'No Label')
        desc = item.get('Description', '')[:60]
        print(f"{indent}{label}: {desc}")

        if item.get('Notes'):
            notes_count = len(item['Notes'])
            print(f"{indent}  [Has {notes_count} note(s)]")

        if item.get('SubStructures'):
            print_hierarchy_sample(item['SubStructures'], depth + 1, max_depth)


def main():
    docx_file = Path(r"c:\Users\Macla\Desktop\AI\DSGL Docs\F2024L01024.docx")
    output_file = Path(r"c:\Users\Macla\Desktop\AI\dual_use_list_parsed.json")

    if not docx_file.exists():
        print(f"Error: File not found: {docx_file}")
        return

    print("Dual-Use List Parser")
    print("=" * 80)

    parser = DualUseParser()
    hierarchy = parser.parse_document(str(docx_file))

    if not hierarchy:
        print("Error: No data parsed")
        return

    # Save to JSON
    print(f"\nSaving to: {output_file}")
    with open(output_file, 'w', encoding='utf-8') as f:
        json.dump(hierarchy, f, indent=2, ensure_ascii=False)

    print(f"Successfully saved!")

    # Print statistics
    total_items = count_items(hierarchy)
    print(f"\nStatistics:")
    print(f"  Top-level items: {len(hierarchy)}")
    print(f"  Total items (including nested): {total_items}")

    # Print sample
    print(f"\nSample hierarchy (first 5 items at each level, max depth 3):")
    print("-" * 80)
    print_hierarchy_sample(hierarchy)

    return hierarchy


if __name__ == "__main__":
    main()
